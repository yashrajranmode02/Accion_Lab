"""
Document Ingestion Module for FAQ RAG Chatbot.

This module handles loading FAQ documents from various file formats (md, txt, docx).
It scans the raw_docs directory and returns Document objects containing text and metadata.
"""

from pathlib import Path
from typing import List, Dict
import docx


class Document:
    """
    Represents a loaded document with text content and metadata.
    
    Attributes:
        text (str): The complete text content of the document.
        source (str): The source filename of the document.
        
    Libraries:
        None (standard Python dataclass-like class)
    """
    
    def __init__(self, text: str, source: str):
        """
        Initialize a Document object.
        
        Args:
            text (str): The text content of the document
            source (str): The source filename
            
        Returns:
            None
            
        Libraries:
            None
        """
        self.text = text
        self.source = source
    
    def __repr__(self):
        return f"Document(source='{self.source}', length={len(self.text)})"


def load_markdown(file_path: Path) -> str:
    """
    Load text from a markdown file.
    
    Args:
        file_path (Path): Path to the markdown file
        
    Returns:
        str: The text content of the file
        
    Libraries:
        pathlib
    """
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()


def load_text(file_path: Path) -> str:
    """
    Load text from a plain text file.
    
    Args:
        file_path (Path): Path to the text file
        
    Returns:
        str: The text content of the file
        
    Libraries:
        pathlib
    """
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()


def load_docx(file_path: Path) -> str:
    """
    Load text from a docx file.
    
    Args:
        file_path (Path): Path to the docx file
        
    Returns:
        str: The text content extracted from all paragraphs
        
    Libraries:
        python-docx
    """
    doc = docx.Document(file_path)
    paragraphs = [para.text for para in doc.paragraphs]
    return '\n'.join(paragraphs)


def load_documents(raw_docs_dir: Path, supported_formats: List[str]) -> List[Document]:
    """
    Load all supported documents from the raw_docs directory.
    
    Scans the directory for files with supported extensions and loads their content
    into Document objects with appropriate metadata.
    
    Args:
        raw_docs_dir (Path): Path to directory containing raw documents
        supported_formats (List[str]): List of supported file extensions (e.g., ['.md', '.txt', '.docx'])
        
    Returns:
        List[Document]: List of Document objects containing text and metadata
        
    Libraries:
        pathlib, python-docx
    """
    documents = []
    
    if not raw_docs_dir.exists():
        raise FileNotFoundError(f"Raw docs directory not found: {raw_docs_dir}")
    
    # Scan directory for supported files
    for file_path in raw_docs_dir.iterdir():
        if file_path.is_file() and file_path.suffix.lower() in supported_formats:
            try:
                # Load based on file extension
                if file_path.suffix.lower() == '.md':
                    text = load_markdown(file_path)
                elif file_path.suffix.lower() == '.txt':
                    text = load_text(file_path)
                elif file_path.suffix.lower() == '.docx':
                    text = load_docx(file_path)
                else:
                    continue
                
                # Create Document object
                doc = Document(text=text, source=file_path.name)
                documents.append(doc)
                print(f"✓ Loaded: {file_path.name} ({len(text)} chars)")
                
            except Exception as e:
                print(f"✗ Error loading {file_path.name}: {e}")
                continue
    
    print(f"\nTotal documents loaded: {len(documents)}")
    return documents


if __name__ == "__main__":
    """
    Example usage of the ingestion module.
    """
    from config import RAW_DOCS_DIR, SUPPORTED_FORMATS
    
    docs = load_documents(RAW_DOCS_DIR, SUPPORTED_FORMATS)
    for doc in docs:
        print(doc)
